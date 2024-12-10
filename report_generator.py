#### 
import datetime
import json
import os
import string
import tempfile
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

class RelatorioAuditoria:

    textos_parametrizados = {}
    if os.path.exists('parametrizacoes/parametrizacao_relatorio.json'):
        with open('parametrizacoes/parametrizacao_relatorio.json', 'r', encoding='utf-8') as f:
            textos_parametrizados = json.load(f)
    else:
        # Fallback caso não exista o arquivo
        textos_parametrizados = {
            "texto_fixo_da_auditoria": "Texto fixo da auditoria.",
            "texto_fixo_do_planejamento": "Texto fixo do planejamento.",
            "texto_fixo_da_matriz_achados": "Texto fixo da matriz de achados."
        }

    texto_fixo_da_auditoria = textos_parametrizados.get('texto_fixo_da_auditoria', "")
    texto_fixo_do_planejamento = textos_parametrizados.get('texto_fixo_do_planejamento', "")
    texto_fixo_da_matriz_achados = textos_parametrizados.get('texto_fixo_da_matriz_achados', "")
    texto_resultado = "Por meio dos procedimentos acima, chegou-se às seguintes conclusões: \n"

    def __init__(self, cnpj:str, cgf:str, ordem_servico:str, proc_sei:str, data_inicio:datetime.datetime, 
                 data_fim:datetime.datetime, nome_arquivo:str, 
                 caminho_logo:str="sefaz_rr_logo.png", caminho_json_planejamento:str="parametrizacoes\criterios_auditoria.json"):
        
        self.cnpj = cnpj
        self.cgf = cgf
        self.ordem_servico = ordem_servico
        self.proc_sei = proc_sei
        self.data_inicio = data_inicio
        self.data_fim = data_fim
        self.nome_arquivo = nome_arquivo
        self.caminho_logo = caminho_logo
        self.caminho_json_planejamento = caminho_json_planejamento
        
        self.dict_planejamento = self._ler_textos_planejamento()
        
        self.lista_procedimentos = []
        
        self.document = None
        

    def _ler_textos_planejamento(self):
        if os.path.exists(self.caminho_json_planejamento):
            with open(self.caminho_json_planejamento, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
        

    def _configurar_estilos_documento(self):
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        

    def _inserir_cabecalho(self):
        if os.path.exists(self.caminho_logo):
            self.document.add_picture(self.caminho_logo, width=Inches(1.5))
        
        titulo_SEFAZ = self.document.add_paragraph("SECRETARIA DE ESTADO DA FAZENDA DE RORAIMA")
        titulo_SEFAZ.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_SEFAZ.runs[0].bold = True

        titulo_DIFIS = self.document.add_paragraph("DIVISÃO DE FISCALIZAÇÃO DE ESTABELECIMENTOS")
        titulo_DIFIS.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_DIFIS.runs[0].bold = True

        self.document.add_paragraph('')
        titulo = self.document.add_paragraph("RELATÓRIO PRELIMINAR DE AUDITORIA")
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo.runs[0].bold = True
        
        self.document.add_paragraph("") # espaço
        
        
    def _criar_capitulo_da_auditoria(self):
        for _ in range(2):
            self.document.add_paragraph("")

        p = self.document.add_paragraph("DA AUDITORIA")
        p.style = self.document.styles['Normal']
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph(self.texto_fixo_da_auditoria)
        
        self.document.add_paragraph("")
        self.document.add_paragraph("INFORMAÇÕES DO PROCESSO").runs[0].bold = True
        tabela_processo = self.document.add_table(rows=4, cols=2)
        tabela_processo.style = 'Table Grid'
        tabela_processo.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        tabela_processo.cell(0,0).text = "Ordem de Serviço"
        tabela_processo.cell(0,1).text = self.ordem_servico
        
        tabela_processo.cell(1,0).text = "Processo SEI"
        tabela_processo.cell(1,1).text = self.proc_sei
        
        tabela_processo.cell(2,0).text = "Data de Relatório"
        tabela_processo.cell(2,1).text = datetime.datetime.now().strftime("%d/%m/%Y")
        
        tabela_processo.cell(3,0).text = "Período de Fiscalização"
        tabela_processo.cell(3,1).text = f"{self.data_inicio.strftime('%d/%m/%Y')} a {self.data_fim.strftime('%d/%m/%Y')}"
        
        self.document.add_paragraph("")
        
        self.document.add_paragraph("INFORMAÇÕES DA EMPRESA").runs[0].bold = True
        tabela_empresa = self.document.add_table(rows=2, cols=2)
        tabela_empresa.style = 'Table Grid'
        tabela_empresa.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        tabela_empresa.cell(0,0).text = "Número CNPJ"
        tabela_empresa.cell(0,1).text = self.cnpj
        
        tabela_empresa.cell(1,0).text = "Número de Cadastro Geral da Fazenda (CGF)"
        tabela_empresa.cell(1,1).text = self.cgf
        
        self.document.add_paragraph("")
        
        self.document.add_page_break()
        
    def _criar_capitulo_do_planejamento(self):
        p = self.document.add_paragraph("DO PLANEJAMENTO")
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph(self.texto_fixo_do_planejamento)
        
        self.p_marcapos_planejamento = self.document.add_paragraph("")

    def _inserir_conteudo_planejamento(self, nome_metodo):
        data = self.dict_planejamento.get(nome_metodo, None)
        if data is None:
            self.document.add_paragraph(f"Texto não encontrado para '{nome_metodo}' no JSON.")
            return
        
        # Insere título da auditoria
        p_sub = self.document.add_paragraph(data.get("titulo", nome_metodo))
        p_sub.runs[0].bold = True
        
        # Descrição introdutória
        desc_intro = data.get("descricao_introdutoria", "")
        if desc_intro:
            self.document.add_paragraph(desc_intro)
        
        # Critérios
        criterios = data.get("criterios", [])
        for i, crit in enumerate(criterios, start=1):
            p_crit = self.document.add_paragraph(f"{i}. {crit.get('titulo', '')}")
            p_crit.runs[0].bold = True
            self.document.add_paragraph(crit.get('descricao', ''))
            
            fundamentacoes = crit.get("fundamentacao", [])
            if fundamentacoes:
                self.document.add_paragraph("Fundamentação Legal:")
                for fnd in fundamentacoes:
                    self.document.add_paragraph(f" - {fnd}")

    def _criar_capitulo_dos_procedimentos_executados(self):
        p = self.document.add_paragraph("DOS PROCEDIMENTOS EXECUTADOS")
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.p_marcapos_procedimentos = self.document.add_paragraph("")

    def _criar_capitulo_da_matriz_achados(self):
        p = self.document.add_paragraph("DA MATRIZ DE ACHADOS")
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph(self.texto_fixo_da_matriz_achados)
        
        # Agora a tabela terá 5 colunas: Procedimento, Critério, Resultado, Conclusão, Providências
        self.tabela_achados = self.document.add_table(rows=1, cols=5)
        self.tabela_achados.style = 'Table Grid'
        hdr_cells = self.tabela_achados.rows[0].cells
        hdr_cells[0].text = "Procedimento"
        hdr_cells[1].text = "Critério"
        hdr_cells[2].text = "Resultado"
        hdr_cells[3].text = "Conclusão"
        hdr_cells[4].text = "Providências"
        
        self.p_marcapos_matriz = self.document.add_paragraph("")

    def _criar_capitulo_conclusao_providencias(self):
        p = self.document.add_paragraph("CONCLUSÃO E PROVIDÊNCIAS")
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.p_marcapos_conclusao = self.document.add_paragraph("")

    def _inserir_rodape(self):
        sections = self.document.sections
        for section in sections:
            footer = section.footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.text = "Relatório Preliminar de Auditoria – Dados sujeitos a Sigilo Fiscal, conforme CTN. O descumprimento do sigilo fiscal é considerado crime. - Página "
            
            run = footer_paragraph.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')

            fldChar3 = OxmlElement('w:t')
            fldChar3.text = "1"

            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run._r.append(fldChar3)
            run._r.append(fldChar4)
        
    def inserir_procedimento(self, nome_metodo:str, texto:str, resultado:str, imagem=None, texto_adicional: str = None, tabela=None, criterio: str = None):
        # Agora permite informar o critério
        proc = {
            "nome_metodo": nome_metodo,
            "texto": texto,
            "texto_adicional": texto_adicional,
            "imagem": imagem,
            "tabela": tabela,
            "resultado": resultado,
            "criterio": criterio
        }
        self.lista_procedimentos.append(proc)
        
        
    def concluir_auditoria(self):
        self.document = Document()
        
        self._configurar_estilos_documento()
        self._inserir_cabecalho()
        self._criar_capitulo_da_auditoria()
        self._criar_capitulo_do_planejamento()

        # Inserir métodos do planejamento (com base no JSON)
        metodos_unicos = []
        for p in self.lista_procedimentos:
            if p["nome_metodo"] not in metodos_unicos:
                metodos_unicos.append(p["nome_metodo"])

        for metodo in metodos_unicos:
            self._inserir_conteudo_planejamento(metodo)

        self.document.add_page_break()
        self._criar_capitulo_dos_procedimentos_executados()
        
        # Agrupar procedimentos por nome_metodo
        procedimentos_por_metodo = {}
        for p in self.lista_procedimentos:
            nm = p['nome_metodo']
            if nm not in procedimentos_por_metodo:
                procedimentos_por_metodo[nm] = []
            procedimentos_por_metodo[nm].append(p)
        
        # Ordenar os métodos
        metodos_ordenados = procedimentos_por_metodo.keys()

        # Inserir procedimentos conforme o novo formato
        for i, metodo in enumerate(metodos_ordenados, start=1):
            roman = self._int_to_roman(i)
            # Subcapítulo com o nome do procedimento
            subt = self.document.add_paragraph(f"{roman}. {metodo}")
            subt.runs[0].bold = True
            
            # Inserir texto do método 
            par_texto = self.document.add_paragraph(procedimentos_por_metodo[metodo][0]['texto'])
            par_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            par_texto.paragraph_format.first_line_indent = Inches(1)

            # Ordenar procedimentos desse método por critério (opcional)
            # Aqui não é obrigatório, mas se desejar, pode-se ordenar por critério.
            # procedimentos_por_metodo[metodo].sort(key=lambda x: x['criterio'] or "")
            
            for j, p in enumerate(procedimentos_por_metodo[metodo]):
                letra = chr(ord('a') + j)
                # Item com o nome do critério
                if p['criterio']:
                    self.document.add_paragraph("")
                    c_par = self.document.add_paragraph(f"{letra}. Procedimentos para conferir critério: {p['criterio']}")
                else:
                    c_par = self.document.add_paragraph(f"{letra}. Procedimentos para conferir critério: [Critério não informado]")

                
                c_par.runs[0].bold = True

                

                # Texto adicional
                if p['texto_adicional'] is not None:
                    ad_par = self.document.add_paragraph(p['texto_adicional'])
                    ad_par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    ad_par.paragraph_format.first_line_indent = Inches(1)

                # Imagem se houver
                if p['imagem'] is not None:
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmpf:
                        p['imagem'].figure.savefig(tmpf.name)
                        self.document.add_picture(tmpf.name, width=Inches(5))
                        os.remove(tmpf.name)

                # Tabela se houver
                if p['tabela'] is not None:
                    df = p['tabela']
                    t = self.document.add_table(rows=df.shape[0]+1, cols=df.shape[1])
                    t.style = 'Table Grid'
                    for col_idx, col_name in enumerate(df.columns):
                        t.cell(0,col_idx).text = str(col_name)
                    for row_idx in range(df.shape[0]):
                        for col_idx in range(df.shape[1]):
                            t.cell(row_idx+1,col_idx).text = str(df.iat[row_idx,col_idx])

                # Resultado
                if p['resultado']:
                    self.document.add_paragraph(self.texto_resultado)
                    resultado_p = self.document.add_paragraph(p['resultado'])
                    resultado_p.runs[0].bold = True
            
            # Realizar quebra de página a cada fim de método 
            self.document.add_page_break()

        self.document.add_page_break()
        self._criar_capitulo_da_matriz_achados()
        
        tabela_achados = self.tabela_achados
        for p in self.lista_procedimentos:
            row_cells = tabela_achados.add_row().cells
            row_cells[0].text = p['nome_metodo']
            row_cells[1].text = p['criterio'] if p['criterio'] else ""
            row_cells[2].text = p['resultado']
            row_cells[3].text = ""
            row_cells[4].text = ""

        self.document.add_page_break()
        self._criar_capitulo_conclusao_providencias()

        achados = []
        for i in range(1, len(self.tabela_achados.rows)):
            procedimento = self.tabela_achados.cell(i,0).text
            criterio = self.tabela_achados.cell(i,1).text
            resultado = self.tabela_achados.cell(i,2).text
            achados.append({"Procedimento": procedimento, "Critério": criterio, "Resultado": resultado})

        prompt = "Resuma os procedimentos executados, conclusões e providências a partir da seguinte matriz de achados:\n"
        for a in achados:
            prompt += f"- {a['Procedimento']} (Critério: {a['Critério']}): {a['Resultado']}\n"

        texto_conclusao = "Este texto foi criado pela LLM local..."
        self.p_marcapos_conclusao.text = texto_conclusao

        self._inserir_rodape()
        

    def salvar(self, tipo:str='docx'):
        if self.document is None:
            raise RuntimeError("Documento não foi criado. Chame 'concluir_auditoria()' primeiro.")
        if tipo == 'docx':
            self.document.save(f"{self.nome_arquivo}.docx")
        elif tipo == 'pdf':
            pass
    
    def _int_to_roman(self, num):
        val = [
            1000, 900, 500, 400,
            100, 90, 50, 40,
            10, 9, 5, 4,
            1
        ]
        syb = [
            "M", "CM", "D", "CD",
            "C", "XC", "L", "XL",
            "X", "IX", "V", "IV",
            "I"
        ]
        roman_num = ''
        i = 0
        while  num > 0:
            for _ in range(num // val[i]):
                roman_num += syb[i]
                num -= val[i]
            i += 1
        return roman_num
####

