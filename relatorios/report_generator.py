####
import datetime
import json
import os
import string
import tempfile
from parametrizacoes.textos_relatorios import TextosRelatorios
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn


class RelatorioAuditoria:

    textos_parametrizados = {}
    if os.path.exists('parametrizacoes/parametrizacao_relatorio.json'):
        with open('parametrizacoes/parametrizacao_relatorio.json', 'r', encoding='utf-8') as f:
            textos_parametrizados = json.load(f)
    else:
        # Fallback caso não exista o arquivo
        textos_parametrizados = {
            "texto_fixo_da_auditoria": "Texto fixo da auditoria.",
            "texto_fixo_do_planejamento": "Texto fixo do planejamento.",
            "texto_fixo_da_matriz_achados": "Texto fixo da matriz de achados."
        }

    texto_fixo_da_auditoria = textos_parametrizados.get('texto_fixo_da_auditoria', "")
    texto_fixo_do_planejamento = textos_parametrizados.get('texto_fixo_do_planejamento', "")
    texto_fixo_da_matriz_achados = textos_parametrizados.get('texto_fixo_da_matriz_achados', "")
    texto_resultado = "Por meio dos procedimentos acima, chegou-se às seguintes conclusões: \n"

    def __init__(self, cnpj:str, cgf:str, ordem_servico:str, proc_sei:str, data_inicio:datetime.datetime, 
                 data_fim:datetime.datetime, nome_arquivo:str, 
                 caminho_logo:str="sefaz_rr_logo.png", caminho_json_planejamento:str="parametrizacoes\criterios_auditoria.json"):
        
        self.cnpj = cnpj
        self.cgf = cgf
        self.ordem_servico = ordem_servico
        self.proc_sei = proc_sei
        self.data_inicio = data_inicio
        self.data_fim = data_fim
        self.nome_arquivo = nome_arquivo
        self.caminho_logo = caminho_logo
        self.caminho_json_planejamento = caminho_json_planejamento
        
        self.dict_planejamento = self._ler_textos_planejamento()
        
        self.lista_procedimentos = []
        
        self.document = None
        

    def _ler_textos_planejamento(self):
        if os.path.exists(self.caminho_json_planejamento):
            with open(self.caminho_json_planejamento, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
        
    def _configurar_estilos_documento(self):
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

    def _adicionar_titulo(self, texto, nivel=1):
        """
        Adiciona um título com o estilo de Heading apropriado para permitir a geração do sumário no Word.
        nivel=1 => Heading 1 (capítulo)
        nivel=2 => Heading 2 (subcapítulo)
        """
        # Cria o parágrafo
        p = self.document.add_paragraph(texto)
        # Define o estilo de título conforme o nível
        if nivel == 1:
            p.style = self.document.styles['Heading 1']
        elif nivel == 2:
            p.style = self.document.styles['Heading 2']
        else:
            # Caso seja necessário mais níveis, pode-se adicionar logicamente
            p.style = self.document.styles['Heading 1']
        
        # Alinha no centro
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def _inserir_cabecalho(self):
        if os.path.exists(self.caminho_logo):
            self.document.add_picture(self.caminho_logo, width=Inches(1.5))
        
        titulo_SEFAZ = self.document.add_paragraph("SECRETARIA DE ESTADO DA FAZENDA DE RORAIMA")
        titulo_SEFAZ.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_SEFAZ.runs[0].bold = True

        titulo_DIFIS = self.document.add_paragraph("DIVISÃO DE FISCALIZAÇÃO DE ESTABELECIMENTOS")
        titulo_DIFIS.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_DIFIS.runs[0].bold = True

        for _ in range(5):
            self.document.add_paragraph('')
        titulo = self.document.add_paragraph("RELATÓRIO PRELIMINAR DE AUDITORIA")
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo.runs[0].bold = True
        titulo.runs[0].font.size = Pt(20)
        
        self.document.add_paragraph("") # espaço

    # 1. Capa e Folha de Rosto
    def _criar_capitulo_capa_folha_rosto(self):
        # Não precisa ser um heading porque é a capa, mas caso queira, pode usar Heading 1
        # Porém normalmente a capa não entra no sumário.
        p1 = self.document.add_paragraph("Ordem de Serviço: " + self.ordem_servico)
        p1.runs[0].font.size = Pt(10)
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        p2 = self.document.add_paragraph("Processo SEI: " + self.proc_sei)
        p2.runs[0].font.size = Pt(10)
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p3 = self.document.add_paragraph("Data de Emissão do Relatório: " + datetime.datetime.now().strftime("%d/%m/%Y"))
        p3.runs[0].font.size = Pt(10)
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p4 = self.document.add_paragraph("Período da Auditoria: " + f"{self.data_inicio.strftime('%d/%m/%Y')} a {self.data_fim.strftime('%d/%m/%Y')}")
        p4.runs[0].font.size = Pt(10)
        p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 2. Sumário
    def _criar_capitulo_sumario(self):
        self.document.add_page_break()
        # Cria um título de sumário sem heading (geralmente sumário não aparece no sumário)
        p = self.document.add_paragraph("SUMÁRIO")
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Insere um campo TOC (Tabela de Conteúdo) que poderá ser atualizado no Word
        # O Word atualiza ao abrir ou ao pressionar F9 dentro do Word.
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "TOC \\o \"1-3\" \\h \\z \\u"

        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')

        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_sep)
        run._r.append(fldChar_end)

    # 3. Apresentação (Introdução)
    def _criar_capitulo_apresentacao(self):
        self.document.add_page_break()
        self._adicionar_titulo("3. APRESENTAÇÃO", nivel=1)

        # Utilize texto da auditoria aqui
        for texto in TextosRelatorios.texto_apresentacao:
            p = self.document.add_paragraph(texto)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self.document.add_paragraph("")

        # Informações da empresa
        p_info = self.document.add_paragraph("INFORMAÇÕES DA EMPRESA AUDITADA")
        p_info.runs[0].bold = True
        p_info.style = self.document.styles['Heading 2']  # Subtítulo para aparecer no sumário

        tabela_empresa = self.document.add_table(rows=2, cols=2)
        tabela_empresa.style = 'Table Grid'
        tabela_empresa.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        tabela_empresa.cell(0,0).text = "Número CNPJ"
        tabela_empresa.cell(0,1).text = self.cnpj
        
        tabela_empresa.cell(1,0).text = "Número de Cadastro Geral da Fazenda (CGF)"
        tabela_empresa.cell(1,1).text = self.cgf

    # 4. CRITÉRIOS DE AUDITORIA
    def _criar_capitulo_criterios_auditoria(self):
        self.document.add_page_break()
        self._adicionar_titulo("4. CRITÉRIOS DE AUDITORIA", nivel=1)

        # Inserir conteúdos do planejamento (critérios) a partir do JSON
        metodos_unicos = []
        for p in self.lista_procedimentos:
            if p["nome_metodo"] not in metodos_unicos:
                metodos_unicos.append(p["nome_metodo"])

        for metodo in metodos_unicos:
            # Subtítulo para cada método
            p_sub = self._adicionar_titulo(metodo, nivel=2)
            self._inserir_conteudo_planejamento(metodo)

    def _inserir_conteudo_planejamento(self, nome_metodo):
        data = self.dict_planejamento.get(nome_metodo, None)
        if data is None:
            self.document.add_paragraph(f"Texto não encontrado para '{nome_metodo}' no JSON.")
            return
        
        # Descrição introdutória
        desc_intro = data.get("descricao_introdutoria", "")
        if desc_intro:
            self.document.add_paragraph(desc_intro)
        
        # Critérios
        criterios = data.get("criterios", [])
        for i, crit in enumerate(criterios, start=1):
            p_crit = self.document.add_paragraph(f"{i}. {crit.get('titulo', '')}")
            self.document.add_paragraph("")
            p_crit.runs[0].bold = True
            desc = self.document.add_paragraph(crit.get('descricao', ''))
            desc.paragraph_format.first_line_indent = Inches(0.5)
            self.document.add_paragraph("")
            
            fundamentacoes = crit.get("fundamentacao", [])
            if fundamentacoes:
                fund_legal = self.document.add_paragraph("Fundamentação Legal:\n")
                fund_legal.runs[0].bold = True
                for fundamentacao in fundamentacoes:
                    self.document.add_paragraph(f" - {fundamentacao.get('referencia', '')}")
                    p = self.document.add_paragraph(fundamentacao.get("comentario", ''))
                    self.document.add_paragraph("")
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Inches(0.5)

    # Metodologia e Procedimentos de Auditoria
    def _criar_capitulo_metodologia(self):
        self.document.add_page_break()
        self._adicionar_titulo("5. PROCEDIMENTOS DE AUDITORIA", nivel=1)
        
        self.document.add_paragraph("A seguir, apresentam-se os procedimentos adotados na execução da auditoria, incluindo análise documental, testes de controle, cruzamento de dados e outras metodologias empregadas:")

        # Agrupar procedimentos por método
        procedimentos_por_metodo = {}
        for pp in self.lista_procedimentos:
            nm = pp['nome_metodo']
            if nm not in procedimentos_por_metodo:
                procedimentos_por_metodo[nm] = []
            procedimentos_por_metodo[nm].append(pp)
        
        metodos_ordenados = procedimentos_por_metodo.keys()
        for i, metodo in enumerate(metodos_ordenados, start=1):
            roman = self._int_to_roman(i)
            # Cada método como subcapítulo
            subt = self._adicionar_titulo(f"{roman}. {metodo}", nivel=2)
            
            # Texto do primeiro procedimento desse método (introdução)
            par_texto = self.document.add_paragraph(procedimentos_por_metodo[metodo][0]['texto'])
            par_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            par_texto.paragraph_format.first_line_indent = Inches(0.5)

            for j, ppp in enumerate(procedimentos_por_metodo[metodo]):
                letra = chr(ord('a') + j)
                if ppp['criterio']:
                    c_par = self.document.add_paragraph(f"{letra}. Procedimentos para conferir critério: {ppp['criterio']}")
                else:
                    c_par = self.document.add_paragraph(f"{letra}. Procedimentos para conferir critério: [Critério não informado]")
                
                c_par.runs[0].bold = True

                if ppp['texto_adicional'] is not None:
                    ad_par = self.document.add_paragraph(ppp['texto_adicional'])
                    ad_par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    ad_par.paragraph_format.first_line_indent = Inches(0.5)

                if ppp['imagem'] is not None:
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmpf:
                        ppp['imagem'].figure.savefig(tmpf.name)
                        self.document.add_picture(tmpf.name, width=Inches(0.5))
                        os.remove(tmpf.name)

                if ppp['tabela'] is not None:
                    df = ppp['tabela']
                    t = self.document.add_table(rows=df.shape[0]+1, cols=df.shape[1])
                    t.style = 'Table Grid'
                    for col_idx, col_name in enumerate(df.columns):
                        t.cell(0,col_idx).text = str(col_name)
                    for row_idx in range(df.shape[0]):
                        for col_idx in range(df.shape[1]):
                            t.cell(row_idx+1,col_idx).text = str(df.iat[row_idx,col_idx])

                if ppp['resultado']:
                    self.document.add_paragraph(self.texto_resultado)
                    resultado_p = self.document.add_paragraph(ppp['resultado'])
                    resultado_p.runs[0].bold = True
                    resultado_p.aligment = WD_ALIGN_PARAGRAPH.CENTER
                    self.document.add_paragraph("")

    # 6. Sumário das Constatações
    def _criar_capitulo_sumario_constatacoes(self):
        self.document.add_page_break()
        self._adicionar_titulo("6. SUMÁRIO DAS CONSTATAÇÕES", nivel=1)

        self.document.add_paragraph(self.texto_fixo_da_matriz_achados)
        
        # Matriz de achados (sumário)
        self.tabela_achados = self.document.add_table(rows=1, cols=5)
        self.tabela_achados.style = 'Table Grid'
        hdr_cells = self.tabela_achados.rows[0].cells
        hdr_cells[0].text = "Número"
        hdr_cells[0].width = Inches(0.5)
        hdr_cells[1].text = "Procedimento"
        hdr_cells[2].text = "Critério"
        hdr_cells[3].text = "Resultado"
        hdr_cells[4].text = "Evidência de Auditoria"

        for i, p in enumerate(self.lista_procedimentos, start=1):
            row_cells = self.tabela_achados.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = p['nome_metodo']
            row_cells[2].text = p['criterio'] if p['criterio'] else ""
            row_cells[3].text = p['resultado']
            row_cells[4].text = p['anexo']

        for paragraph in hdr_cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 7. Análise Detalhada das Constatações
    def _criar_capitulo_analise_constatacoes(self):
        self.document.add_page_break()
        self._adicionar_titulo("7. ANÁLISE DETALHADA DAS CONSTATAÇÕES", nivel=1)

        p2 = self.document.add_paragraph(TextosRelatorios.texto_analise_achados)
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Matriz de conclusões
        self.tabela_achados = self.document.add_table(rows=1, cols=4)
        self.tabela_achados.style = 'Table Grid'
        hdr_cells = self.tabela_achados.rows[0].cells
        hdr_cells[0].text = "Número"
        hdr_cells[0].width = Inches(0.5)
        hdr_cells[1].text = "Critério"
        hdr_cells[2].text = "Resultado"
        hdr_cells[3].text = "Conclusões"


        for i, p in enumerate(self.lista_procedimentos, start=1):
            row_cells = self.tabela_achados.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = p['criterio'] if p['criterio'] else ""
            row_cells[2].text = p['resultado']
            row_cells[3].text = ""

        for paragraph in hdr_cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 8. Conclusões e Opiniões do Auditor Fiscal
    def _criar_capitulo_conclusoes(self):
        self.document.add_page_break()
        self._adicionar_titulo("8. CONCLUSÃO E OPINIÃO DO AUDITOR FISCAL", nivel=1)

        achados = []
        for i in range(1, len(self.tabela_achados.rows)):
            procedimento = self.tabela_achados.cell(i,0).text
            criterio = self.tabela_achados.cell(i,1).text
            resultado = self.tabela_achados.cell(i,2).text
            achados.append({"Procedimento": procedimento, "Critério": criterio, "Resultado": resultado})

        prompt = "Resuma os procedimentos executados, conclusões e providências a partir da seguinte matriz de achados:\n"
        for a in achados:
            prompt += f"- {a['Procedimento']} (Critério: {a['Critério']}): {a['Resultado']}\n"

        texto_conclusao = "Com base nas constatações, conclui-se que o contribuinte apresenta inconsistências em determinados critérios, conforme evidenciado nos resultados. A opinião do auditor fiscal é de que sejam adotadas medidas corretivas e, caso necessário, lançamentos de ofício, bem como outras providências legais."
        self.document.add_paragraph(texto_conclusao)


    # 9. Recomendações e Determinações
    def _criar_capitulo_recomendacoes(self):
        self.document.add_page_break()
        self._adicionar_titulo("9. RECOMENDAÇÕES E DETERMINAÇÕES", nivel=1)
        self.document.add_paragraph("")


    def _inserir_rodape(self):
        sections = self.document.sections
        for section in sections:
            footer = section.footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.text = "Relatório Preliminar de Auditoria – Dados sujeitos a Sigilo Fiscal, conforme CTN. O descumprimento do sigilo fiscal é considerado crime. - Página "
            
            run = footer_paragraph.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')

            fldChar3 = OxmlElement('w:t')
            fldChar3.text = "1"

            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run._r.append(fldChar3)
            run._r.append(fldChar4)
        
    def inserir_procedimento(self, nome_metodo:str, texto:str, resultado:str, anexo: str,imagem=None, texto_adicional: str = None, tabela=None, criterio: str = None):
        proc = {
            "nome_metodo": nome_metodo,
            "texto": texto,
            "texto_adicional": texto_adicional,
            "imagem": imagem,
            "tabela": tabela,
            "resultado": resultado,
            "criterio": criterio,
            "anexo": anexo
        }
        self.lista_procedimentos.append(proc)
        
        
    def concluir_auditoria(self):
        self.document = Document()
        
        self._configurar_estilos_documento()
        self._inserir_cabecalho()
        
        # Chamar os capítulos na ordem solicitada
        self._criar_capitulo_capa_folha_rosto()          # 1
        self._criar_capitulo_sumario()                   # 2
        self._criar_capitulo_apresentacao()              # 3
        self._criar_capitulo_criterios_auditoria()       # 4
        self._criar_capitulo_metodologia()               # 5
        self._criar_capitulo_sumario_constatacoes()      # 6
        self._criar_capitulo_analise_constatacoes()      # 7
        self._criar_capitulo_conclusoes()                # 8
        self._criar_capitulo_recomendacoes()             # 9

        self._inserir_rodape()

    def salvar(self, tipo:str='docx'):
        if self.document is None:
            raise RuntimeError("Documento não foi criado. Chame 'concluir_auditoria()' primeiro.")
        if tipo == 'docx':
            self.document.save(f"{self.nome_arquivo}.docx")
        elif tipo == 'pdf':
            pass
    
    def _int_to_roman(self, num):
        val = [
            1000, 900, 500, 400,
            100, 90, 50, 40,
            10, 9, 5, 4,
            1
        ]
        syb = [
            "M", "CM", "D", "CD",
            "C", "XC", "L", "XL",
            "X", "IX", "V", "IV",
            "I"
        ]
        roman_num = ''
        i = 0
        while  num > 0:
            for _ in range(num // val[i]):
                roman_num += syb[i]
                num -= val[i]
            i += 1
        return roman_num
####
